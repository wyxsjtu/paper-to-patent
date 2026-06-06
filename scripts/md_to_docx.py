#!/usr/bin/env python3
"""
Convert patent disclosure Markdown to Word (.docx).

Conversion strategy (in order):
  1. pandoc -- best quality, handles LaTeX math, tables, images
  2. python-docx -- basic fallback if pandoc not installed

Usage:
  python3 md_to_docx.py <input.md> [output.docx]

The script looks for PNG images referenced in the Markdown
in the same directory as the input file.
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
import urllib.request
from pathlib import Path

CHINESE_BODY_FONT = "宋体"
CHINESE_HEADING_FONT = "黑体"
WESTERN_BODY_FONT = "Times New Roman"


# ---------------------------------------------------------------------------
# Method 1: pandoc
# ---------------------------------------------------------------------------

def _patch_reference_docx(ref_path: Path) -> None:
    """
    Open reference.docx and overwrite CJK font settings in key styles so that
    pandoc output uses 宋体/黑体 instead of the default MS Gothic.
    Preserves all other pandoc-required styles already in the file.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        return

    try:
        doc = Document(str(ref_path))
    except Exception as e:
        print(f"[pandoc] Cannot open reference.docx for font patch: {e}", file=sys.stderr)
        return

    def _set_cjk_fonts(style, east_asia: str, western: str, size_pt: float, bold=None):
        try:
            rpr = style.element.get_or_add_rPr()
            rFonts = rpr.get_or_add_rFonts()
            # Set explicit font names
            rFonts.set(qn("w:eastAsia"), east_asia)
            rFonts.set(qn("w:ascii"), western)
            rFonts.set(qn("w:hAnsi"), western)
            rFonts.set(qn("w:cs"), western)
            # Remove *Theme attributes: in OOXML, w:eastAsiaTheme overrides w:eastAsia
            # so MS Gothic (or any theme CJK font) would take priority unless removed
            for attr in (qn("w:eastAsiaTheme"), qn("w:asciiTheme"),
                         qn("w:hAnsiTheme"), qn("w:cstheme")):
                rFonts.attrib.pop(attr, None)
            style.font.name = western
            style.font.size = Pt(size_pt)
            if bold is not None:
                style.font.bold = bold
        except Exception:
            pass

    style_configs = [
        ("Normal",                  CHINESE_BODY_FONT,    WESTERN_BODY_FONT, 12,   None),
        ("Default Paragraph Font",  CHINESE_BODY_FONT,    WESTERN_BODY_FONT, 12,   None),
        ("Title",                   CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 16,   True),
        ("Heading 1",               CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 14,   True),
        ("Heading 2",               CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 13,   True),
        ("Heading 3",               CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 12,   True),
        ("Heading 4",               CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 12,   True),
        ("List Bullet",             CHINESE_BODY_FONT,    WESTERN_BODY_FONT, 12,   None),
        ("List Number",             CHINESE_BODY_FONT,    WESTERN_BODY_FONT, 12,   None),
    ]

    from docx.oxml import OxmlElement

    def _reset_color_to_auto(style):
        """Remove any explicit color from the style's rPr so text renders as black."""
        try:
            rpr = style.element.get_or_add_rPr()
            # Remove existing w:color elements (could be blue from pandoc defaults)
            for color_el in rpr.findall(qn("w:color")):
                rpr.remove(color_el)
            # Set w:color val="auto" (inherits theme/black)
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), "auto")
            rpr.append(color_el)
        except Exception:
            pass

    for name, ea, western, size, bold in style_configs:
        try:
            s = doc.styles[name]
            _set_cjk_fonts(s, ea, western, size, bold)
            _reset_color_to_auto(s)
        except KeyError:
            pass  # style absent in this reference.docx — skip

    try:
        doc.save(str(ref_path))
    except Exception as e:
        print(f"[pandoc] Cannot save patched reference.docx: {e}", file=sys.stderr)


def convert_with_pandoc(md_path: str, docx_path: str) -> bool:
    pandoc_bin = _find_pandoc()
    if not pandoc_bin:
        print("[pandoc] Not found in PATH", file=sys.stderr)
        return False

    ref_docx = Path(__file__).parent.parent / "templates" / "reference.docx"

    if ref_docx.is_file():
        _patch_reference_docx(ref_docx)

    cmd = [
        pandoc_bin,
        md_path,
        "-o", docx_path,
        "--wrap=none",
        # LaTeX math
        "--mathml",
    ]

    if ref_docx.is_file():  # re-check after patch (save could theoretically fail)
        cmd += ["--reference-doc", str(ref_docx)]

    # Set resource path so images are found relative to the md file
    cmd += ["--resource-path", str(Path(md_path).parent)]

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
            cwd=str(Path(md_path).parent),
        )
        if result.returncode != 0:
            print(f"[pandoc] stderr: {result.stderr[:500]}", file=sys.stderr)
            return False
        return Path(docx_path).is_file()
    except subprocess.TimeoutExpired:
        print("[pandoc] Timed out", file=sys.stderr)
        return False
    except Exception as exc:
        print(f"[pandoc] {exc}", file=sys.stderr)
        return False


def _find_pandoc() -> str:
    pandoc_bin = shutil.which("pandoc")
    if pandoc_bin:
        return pandoc_bin
    try:
        import pypandoc
        candidate = pypandoc.get_pandoc_path()
        if candidate and Path(candidate).is_file():
            return candidate
    except Exception:
        pass
    return ""


# ---------------------------------------------------------------------------
# LaTeX → OMML (Office Math Markup Language) for Word native equations
# ---------------------------------------------------------------------------

# Multiple candidate URLs for the MML2OMML.XSL stylesheet (Microsoft standard)
_MML2OMML_URLS = [
    "https://raw.githubusercontent.com/jmcnamara/excel-writer-xlsx/master/t/utils/MML2OMML.XSL",
    "https://raw.githubusercontent.com/nicowillis/MML2OMML/master/MML2OMML.XSL",
]


def _get_mml2omml_xsl() -> str:
    """Return MML2OMML.XSL content; download and cache if not present."""
    xsl_path = Path(__file__).parent.parent / "templates" / "MML2OMML.XSL"
    if xsl_path.is_file() and xsl_path.stat().st_size > 1000:
        return xsl_path.read_text(encoding="utf-8")
    for url in _MML2OMML_URLS:
        try:
            with urllib.request.urlopen(url, timeout=15) as resp:
                content = resp.read().decode("utf-8")
            if len(content) > 1000:
                xsl_path.write_text(content, encoding="utf-8")
                print(f"[latex2omml] Downloaded MML2OMML.XSL -> {xsl_path}", file=sys.stderr)
                return content
        except Exception as e:
            print(f"[latex2omml] Cannot fetch {url}: {e}", file=sys.stderr)
    return ""


def _ensure_latex2mathml():
    """Import latex2mathml, installing it if necessary. Returns the module or None."""
    try:
        import latex2mathml.converter
        return latex2mathml.converter
    except ImportError:
        try:
            print("[latex2omml] Installing latex2mathml...", file=sys.stderr)
            subprocess.run(
                [sys.executable, "-m", "pip", "install", "latex2mathml", "-q"],
                check=True, capture_output=True,
            )
            import latex2mathml.converter
            return latex2mathml.converter
        except Exception as e:
            print(f"[latex2omml] Cannot install latex2mathml: {e}", file=sys.stderr)
            return None


def _latex_to_omml(latex_str: str) -> str:
    """
    Convert a LaTeX math string to OMML XML for insertion into Word .docx.
    Returns an OMML XML string, or empty string on failure.
    Conversion chain: LaTeX → MathML (latex2mathml) → OMML (MML2OMML.XSL via lxml).
    """
    converter = _ensure_latex2mathml()
    if converter is None:
        return ""

    # latex2mathml expects the raw expression without surrounding $$ delimiters
    latex_clean = latex_str.strip().lstrip("$").rstrip("$").strip()
    try:
        mathml = converter.convert(latex_clean)
    except Exception as e:
        print(f"[latex2omml] latex2mathml failed: {e}", file=sys.stderr)
        return ""

    xsl_str = _get_mml2omml_xsl()
    if not xsl_str:
        return ""

    try:
        from lxml import etree
        xsl_tree = etree.fromstring(xsl_str.encode("utf-8"))
        transform = etree.XSLT(xsl_tree)
        mathml_tree = etree.fromstring(mathml.encode("utf-8"))
        omml_tree = transform(mathml_tree)
        return etree.tostring(omml_tree, encoding="unicode")
    except Exception as e:
        print(f"[latex2omml] XSLT transform failed: {e}", file=sys.stderr)
        return ""


def _insert_omml_paragraph(doc, latex_str: str, center: bool = True) -> bool:
    """
    Add a block-level OMML equation paragraph to the document.
    Returns True on success, False on fallback.
    """
    omml = _latex_to_omml(latex_str)
    if not omml:
        return False
    try:
        from docx.oxml import parse_xml
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        # Insert OMML element into paragraph XML
        omml_el = parse_xml(omml)
        p._element.append(omml_el)
        return True
    except Exception as e:
        print(f"[latex2omml] docx insert failed: {e}", file=sys.stderr)
        return False


# ---------------------------------------------------------------------------
# Method 2: python-docx fallback
# ---------------------------------------------------------------------------

def convert_with_python_docx(md_path: str, docx_path: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[python-docx] Not installed. Run: pip install python-docx", file=sys.stderr)
        return False

    try:
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
    except Exception as exc:
        print(f"[python-docx] Cannot read {md_path}: {exc}", file=sys.stderr)
        return False

    doc = Document()
    _setup_docx_styles(doc)
    md_dir = Path(md_path).parent

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = _strip_inline_md(m.group(2))
            style = f"Heading {level}"
            try:
                doc.add_heading(text, level=level)
            except Exception:
                doc.add_paragraph(text, style="Normal")
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Images  ![alt](path)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            alt, img_path = m.group(1), m.group(2)
            full_img = md_dir / img_path
            if full_img.is_file():
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(full_img), width=Inches(5.5))
                    cap = doc.add_paragraph(alt, style="Caption")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as exc:
                    print(f"[python-docx] Image {full_img}: {exc}", file=sys.stderr)
                    doc.add_paragraph(f"[图: {alt}]")
            else:
                doc.add_paragraph(f"[图: {alt} (file not found: {img_path})]")
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[-| :]+\|?$", lines[i + 1].strip()):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j]:
                table_lines.append(lines[j].rstrip("\n"))
                j += 1
            _add_table_to_docx(doc, table_lines)
            i = j
            continue

        # Block math  $$...$$
        if line.strip().startswith("$$"):
            math_lines = []
            # Collect everything between the opening $$ and closing $$
            inner = line.strip()[2:]  # content after opening $$
            if "$$" in inner:
                # Single-line: $$ expr $$
                latex_str = inner[:inner.index("$$")].strip()
            else:
                math_lines.append(inner)
                i += 1
                while i < len(lines):
                    cur = lines[i].rstrip("\n")
                    if "$$" in cur:
                        math_lines.append(cur[:cur.index("$$")])
                        i += 1
                        break
                    math_lines.append(cur)
                    i += 1
                latex_str = "\n".join(math_lines).strip()
            if latex_str and _insert_omml_paragraph(doc, latex_str):
                continue
            # Fallback: italic text with $$ delimiters
            fallback = f"$${latex_str}$$"
            p = doc.add_paragraph(fallback, style="Normal")
            if p.runs:
                p.runs[0].italic = True
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", line):
            text = _strip_inline_md(re.sub(r"^[-*+]\s+", "", line))
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+[.)]\s+", line):
            text = _strip_inline_md(re.sub(r"^\d+[.)]\s+", "", line))
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph (with inline math support)
        _add_paragraph_with_inline_math(doc, line)
        i += 1

    try:
        doc.save(docx_path)
        return True
    except Exception as exc:
        print(f"[python-docx] Save failed: {exc}", file=sys.stderr)
        return False


def _setup_docx_styles(doc):
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE

    def set_font(style, east_asia: str, western: str, size_pt: float, bold=None):
        style.font.name = western
        style.font.size = Pt(size_pt)
        if bold is not None:
            style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style._element.rPr.rFonts.set(qn("w:ascii"), western)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), western)

    try:
        normal = doc.styles["Normal"]
        set_font(normal, CHINESE_BODY_FONT, WESTERN_BODY_FONT, 12)
        normal.paragraph_format.first_line_indent = Pt(24)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(0)
    except Exception:
        pass

    for i in range(1, 5):
        try:
            h = doc.styles[f"Heading {i}"]
            set_font(h, CHINESE_HEADING_FONT, WESTERN_BODY_FONT, 14 if i == 1 else 12, True)
            h.paragraph_format.first_line_indent = Pt(0)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(6)
        except Exception:
            pass

    for name in ("List Bullet", "List Number", "Caption"):
        try:
            style = doc.styles[name]
            set_font(style, CHINESE_BODY_FONT, WESTERN_BODY_FONT, 12 if name != "Caption" else 10.5)
            style.paragraph_format.first_line_indent = Pt(0)
        except Exception:
            pass


def _strip_inline_md(text: str) -> str:
    """Remove inline Markdown formatting."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def _add_paragraph_with_inline_math(doc, line: str):
    """
    Add a paragraph that may contain inline $...$ formulas.
    Each formula segment becomes an OMML run; plain text becomes normal runs.
    Falls back to stripped plain text if OMML conversion fails.
    """
    from docx.oxml import parse_xml

    # Split on inline $...$ patterns (non-greedy, single-line)
    segments = re.split(r"(\$[^$\n]+?\$)", line)
    if len(segments) == 1:
        # No inline math – plain paragraph
        doc.add_paragraph(_strip_inline_md(line), style="Normal")
        return

    p = doc.add_paragraph(style="Normal")
    for seg in segments:
        if re.match(r"^\$[^$\n]+?\$$", seg):
            latex = seg[1:-1].strip()
            omml = _latex_to_omml(latex)
            if omml:
                try:
                    omml_el = parse_xml(omml)
                    p._element.append(omml_el)
                    continue
                except Exception:
                    pass
            # Fallback: italic run
            run = p.add_run(seg)
            run.italic = True
        else:
            plain = _strip_inline_md(seg)
            if plain:
                p.add_run(plain)


def _add_table_to_docx(doc, table_lines: list):
    from docx.shared import Pt
    rows = []
    for line in table_lines:
        if re.match(r"^\|?[-| :]+\|?$", line.strip()):
            continue
        cells = [c.strip() for c in re.split(r"\|", line.strip("| \t"))]
        if cells:
            rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < ncols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _strip_inline_md(cell_text)
                if r_idx == 0:
                    cell.paragraphs[0].runs[0].bold = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Convert patent disclosure MD to DOCX")
    parser.add_argument("input_md", help="Input Markdown file")
    parser.add_argument("output_docx", nargs="?",
                        help="Output DOCX file (default: same name as input)")
    args = parser.parse_args()

    md_path = args.input_md
    if not os.path.isfile(md_path):
        print(f"[ERROR] File not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    docx_path = args.output_docx or re.sub(r"\.md$", ".docx", md_path, flags=re.I)
    if docx_path == md_path:
        docx_path = md_path + ".docx"

    print(f"[*] Converting {md_path} -> {docx_path}", file=sys.stderr)

    success = convert_with_pandoc(md_path, docx_path)
    if success:
        print(f"[*] pandoc: OK -> {docx_path}", file=sys.stderr)
        return

    print("[*] Falling back to python-docx...", file=sys.stderr)
    success = convert_with_python_docx(md_path, docx_path)
    if success:
        print(f"[*] python-docx: OK -> {docx_path}", file=sys.stderr)
    else:
        print("[ERROR] Both conversion methods failed.", file=sys.stderr)
        print("Install pandoc:     sudo apt install pandoc", file=sys.stderr)
        print("  OR python-docx:   pip install python-docx", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
